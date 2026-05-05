import pdfplumber
import docx
import pytesseract
from PIL import Image
import os

class PolicyExtractor:
    @staticmethod
    def extract_txt(file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read(), 1 # TXT is usually 1 "page" conceptually or we don't count

    @staticmethod
    def extract_docx(file_path):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text), len(doc.paragraphs) # Paragraph count as proxy if pages not available easily

    @staticmethod
    def extract_pdf(file_path, log_callback=None):
        text_content = []
        page_count = 0
        used_ocr = False

        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(page_text)
                else:
                    # Fallback to OCR if no text found on this page
                    if log_callback:
                        log_callback(f"No text layer found on page {page.page_number}. Attempting OCR...")
                    
                    used_ocr = True
                    # Convert page to image for OCR
                    # pdfplumber page to image
                    img = page.to_image(resolution=300).original
                    ocr_text = pytesseract.image_to_string(img)
                    text_content.append(ocr_text)

        return text_content, page_count, used_ocr

    @classmethod
    def extract(cls, file_path, file_type, log_callback=None):
        if file_type == 'PDF':
            return cls.extract_pdf(file_path, log_callback)
        elif file_type == 'DOCX':
            text, count = cls.extract_docx(file_path)
            return [text], count, False
        elif file_type == 'TXT':
            text, count = cls.extract_txt(file_path)
            return [text], count, False
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
